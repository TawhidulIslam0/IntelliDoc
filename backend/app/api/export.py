import io
import json
import uuid
import os
import boto3
import html2text
import zipfile

from docx import Document
from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from botocore.exceptions import ClientError
from botocore.config import Config
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.orm import Session
from typing import Annotated, Optional
from dotenv import load_dotenv

from app.database import get_db
from app.api.users import get_current_user
from app.models.file import File
from app.models.user import User
from app.models.profile import Profile
from app.models.folder import Folder

load_dotenv()

s3 = boto3.client(
    "s3",
    region_name=os.getenv("AWS_REGION"),
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    config=Config(signature_version="s3v4")
)

BUCKET_NAME = os.getenv("AWS_BUCKET_NAME")

router = APIRouter(prefix="/files", tags=["export"])

SUPPORTED_FORMATS = {"pdf", "docx", "md", "txt"}

# To Markdown
def idoc_to_markdown(pages: list[str]) -> str:
    # html2text converts HTML tags to their Markdown versions. 
    converter = html2text.HTML2Text()
    converter.ignore_links = False
    return "\n\n---\n\n".join(converter.handle(page) for page in pages) # Joins pages

# To TXT
def idoc_to_txt(pages: list[str]) -> str:
    # Extract text from HTML on each page and join with double newlines
    text_content = []
    for page_html in pages:
        soup = BeautifulSoup(page_html, "html.parser")
        text_content.append(soup.get_text(separator="\n"))
    return "\n\n".join(text_content)

# DOCX helper for inline formatting within a single block element. 
def _add_runs(para, element) -> None:
    for child in element.children:
        tag = getattr(child, "name", None)
        text = child.get_text() if hasattr(child, "get_text") else str(child)
        if not text:
            continue
        run = para.add_run(text)
        if tag in ("b", "strong"):
            run.bold = True
        elif tag in ("i", "em"):
            run.italic = True
        elif tag == "u":
            run.underline = True

# To DOCX
def idoc_to_docx(pages: list[str]) -> bytes:
    doc = Document()
    for i, page_html in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        soup = BeautifulSoup(page_html, "html.parser")
        for element in soup.children:
            tag = getattr(element, "name", None)
            text = element.get_text()
            if not text.strip():
                continue
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                doc.add_heading(text, level = int(tag[1]))
            elif tag in ("p", "div"):
                para = doc.add_paragraph()
                _add_runs(para, element)
            elif tag == "ul":
                for li in element.find_all("li", recursive = False):
                    doc.add_paragraph(li.get_text(), style = "List Bullet")
            elif tag == "ol":
                for li in element.find_all("li", recursive = False):
                    doc.add_paragraph(li.get_text(), style = "List Number")
            elif tag is None:  # plain text node
                if text.strip():
                    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# To PDF
def idoc_to_pdf(pages: list[str]) -> bytes:
    page_divs = "".join(f'<div class="page">{page}</div>' for page in pages)
    full_html = (
        "<html><head><style>"
        "body { font-family: Arial, sans-serif; font-size: 12pt; }"
        ".page { page-break-after: always; }"
        ".page:last-child { page-break-after: auto; }"
        "</style></head><body>"
        + page_divs
        + "</body></html>"
    )
    buf = io.BytesIO()
    result = pisa.CreatePDF(full_html, dest = buf)
    if result.err:
        raise ValueError("PDF generation failed")
    buf.seek(0)
    return buf.read()


@router.get("/{file_id}/export")
async def export_file(
    file_id: uuid.UUID,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
    format: str = Query(..., description = "Export format: MD, DOCX, PDF, or TXT"),
    profile_id: Optional[uuid.UUID] = None,
):
    if format not in SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{format}'. Choose from: {', '.join(SUPPORTED_FORMATS)}"
        )

    # Default to the user's default profile if none provided
    if not profile_id:
        profile = db.scalar(
            select(Profile).where(
                Profile.owner_id == current_user.id,
                Profile.is_default == True
            )
        )
        if not profile:
            raise HTTPException(status_code=404, detail="Default profile not found")
        profile_id = profile.id

    file = db.scalar(
        select(File).where(
            File.id == file_id,
            File.owner_id == current_user.id,
            File.profile_id == profile_id
        )
    )

    if not file:
        raise HTTPException(status_code=404, detail="File not found")

    # Fetch idoc content from S3
    try:
        response = s3.get_object(Bucket=BUCKET_NAME, Key=file.s3_key)
        content = json.loads(response["Body"].read().decode("utf-8"))
    except ClientError as e:
        if e.response["Error"]["Code"] == "NoSuchKey":
            raise HTTPException(status_code=404, detail="File content not found in storage")
        raise HTTPException(status_code=500, detail="Failed to fetch file from storage")

    pages = content.get("pages", [""])
    if not isinstance(pages, list):
        raise HTTPException(status_code=422, detail="Invalid idoc structure: 'pages' must be a list")

    base_name = file.name.removesuffix(".idoc")

    try:
        if format == "md":
            data = idoc_to_markdown(pages).encode("utf-8")
            media_type = "text/markdown"
            filename = f"{base_name}.md"
        elif format == "docx":
            data = idoc_to_docx(pages)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{base_name}.docx"
        elif format == "txt":
            data = idoc_to_txt(pages).encode("utf-8")
            media_type = "text/plain"
            filename = f"{base_name}.txt"
        else:  # pdf
            data = idoc_to_pdf(pages)
            media_type = "application/pdf"
            filename = f"{base_name}.pdf"
    except Exception:
        raise HTTPException(status_code=500, detail=f"Failed to convert document to {format}")

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

def _write_folder_to_zip(
    db: Session,
    zip_file: zipfile.ZipFile,
    folder: Folder,
    current_user: User,
    profile_id: uuid.UUID,
    parent_path: str = "",
) -> None:
    folder_path = f"{parent_path}{folder.name}/"

    files = db.scalars(
        select(File).where(
            File.owner_id == current_user.id,
            File.profile_id == profile_id,
            File.folder_id == folder.id
        )
    ).all()

    for file in files:
        try:
            response = s3.get_object(Bucket=BUCKET_NAME, Key=file.s3_key)
            raw_data = response["Body"].read()
        except ClientError:
            continue

        # Keep every file in its original format when downloading a folder.
        zip_file.writestr(f"{folder_path}{file.name}", raw_data)

    subfolders = db.scalars(
        select(Folder).where(
            Folder.owner_id == current_user.id,
            Folder.profile_id == profile_id,
            Folder.parent_id == folder.id
        )
    ).all()

    for subfolder in subfolders:
        _write_folder_to_zip(
            db=db,
            zip_file=zip_file,
            folder=subfolder,
            current_user=current_user,
            profile_id=profile_id,
            parent_path=folder_path,
        )

@router.get("/folders/{folder_id}/export")
async def export_folder(
    folder_id: uuid.UUID,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
    profile_id: Optional[uuid.UUID] = None,
):
    # Default to the user's default profile if none provided
    if not profile_id:
        profile = db.scalar(
            select(Profile).where(
                Profile.owner_id == current_user.id,
                Profile.is_default == True
            )
        )
        if not profile:
            raise HTTPException(status_code=404, detail="Default profile not found")
        profile_id = profile.id

    folder = db.scalar(
        select(Folder).where(
            Folder.id == folder_id,
            Folder.owner_id == current_user.id,
            Folder.profile_id == profile_id
        )
    )

    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")

    zip_buffer = io.BytesIO()

    try:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            _write_folder_to_zip(
                db=db,
                zip_file=zip_file,
                folder=folder,
                current_user=current_user,
                profile_id=profile_id,
            )
    except Exception:
        raise HTTPException(status_code=500, detail="Failed to export folder")

    zip_buffer.seek(0)

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{folder.name}.zip"'}
    )