"""
text_extractor.py
-----------------
Extracts text from various file formats.
"""

from pathlib import Path
from charset_normalizer import from_path
import pymupdf
import docx
import re
from collections import Counter

def _strip_repeating_headers_footers(pages: list[str]) -> list[str]:
    """Remove lines that appear on many pages --- likely running headers/footers.
    Looks at the first and last non-empty line of each page. If a line
    shows up on more than 50% of pages, it's treated as a header/footer
    and removed from every page.
    """
    if len(pages) < 3:
        return pages
 
    counts: Counter[str] = Counter()
    for page in pages:
        lines = [ln.strip() for ln in page.splitlines() if ln.strip()]
        if not lines:
            continue
        counts[lines[0]] += 1
        if len(lines) > 1:
            counts[lines[-1]] += 1
 
    threshold = len(pages) * 0.5
    repeating = {line for line, n in counts.items() if n > threshold}
    if not repeating:
        return pages
 
    cleaned = []
    for page in pages:
        kept = [ln for ln in page.splitlines() if ln.strip() not in repeating]
        cleaned.append("\n".join(kept))
    return cleaned

# "word-\nword" --- a hyphen at end of line immediately followed by word chars.
_SOFT_HYPHEN_RE = re.compile(r"(\w+)-\n(\w+)")
 
# A single \n that is NOT adjacent to another \n --- i.e. a line wrap,
# not part of a paragraph break.
_SINGLE_NEWLINE_RE = re.compile(r"(?<!\n)\n(?!\n)")
 
# Three or more consecutive newlines --- collapse to a normal paragraph break.
_EXTRA_BLANKS_RE = re.compile(r"\n{3,}")
 
# Runs of spaces/tabs (but NOT newlines).
_HORIZONTAL_WS_RE = re.compile(r"[ \t\f\v]+")
 
 
def _rejoin_hyphenation(text: str) -> str:
    """Join words that were split across lines by hyphen.
    'exam-\\nple' -> 'example'
    """
    return _SOFT_HYPHEN_RE.sub(r"\1\2", text)

def clean(text: str) -> str:
    """Clean raw extracted text.
    Pipeline:
      0. Normalize line endings to \n                 (Windows \r\n, old Mac \r)
      1. Remove trailing whitespace on each line
      2. Rejoin words broken by end-of-line hyphens   ('exam-\\nple' -> 'example')
      3. Turn single newlines into spaces             (line wraps, not paragraphs)
      4. Collapse 3+ newlines to exactly 2            (normalize paragraph breaks)
      5. Collapse runs of spaces/tabs to one space
      6. Strip whitespace on each line and overall
    """
    # Normalize line endings to \n. This is important.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Remove trailing whitespace from every line so blank text containing spaces/tab becomes empty and the single newline regex can work properly.
    text = "\n".join(line.rstrip() for line in text.split("\n"))

    # Order matters: hyphenation must run BEFORE we touch newlines, because
    # the pattern depends on seeing the '\n' right after the '-'.
    text = _rejoin_hyphenation(text)
 
    # Single newline => line wrap => space.
    # Double newlines are preserved because the lookaround refuses to match
    # a \n that is adjacent to another \n.
    text = _SINGLE_NEWLINE_RE.sub(" ", text)
 
    # Normalize 3+ newlines down to exactly 2 (one blank line between paragraphs).
    text = _EXTRA_BLANKS_RE.sub("\n\n", text)
 
    # Collapse horizontal whitespace.
    text = _HORIZONTAL_WS_RE.sub(" ", text)
 
    # Trim each line and the whole string.
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()

def extract_pdf(path: str) -> str:
    """Extract text from a PDF.
 
    Pages are extracted individually, running headers/footers are detected
    and stripped, then pages are joined with blank lines between them.
    """
    doc = pymupdf.open(path)
    try:
        # "text" mode preserves reading order for most documents.
        pages = [page.get_text("text") for page in doc]
    finally:
        doc.close()
 
    # If this looks like a scanned PDF with no embedded text, return empty. (NOT FINISHED)
    if sum(len(p.strip()) for p in pages) < 20:
        return ""
 
    pages = _strip_repeating_headers_footers(pages)
    return "\n\n".join(p for p in pages if p.strip())

def extract_docx(path: str) -> str:
    """Extract text from a DOCX file.
    Paragraphs are preserved as paragraphs (separated by 2 new blanklines)
    """

    parts: list[str] = []

    doc = docx.Document(path)

    # Detects paragraphs
    for para in doc.paragraphs:
        text = para.text
        if text:
            parts.append(text)
    
    # Detects tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    
    return "\n\n".join(parts)

def extract_txt(path: str) -> str:
    """Extract text from a TXT file. Don't assume UTF-8 encoding"""

    result = from_path(path).best()
    if result is not None:
        return str(result)
    return Path(path).read_text(encoding="utf-8", errors="replace")


def extract(path: str) -> str:
    """Extract text from a file depending on file type."""

    suff = Path(path).suffix.lower()

    if suff == ".pdf":
        return extract_pdf(path)
    elif suff == ".docx":
        return extract_docx(path)
    elif suff == ".txt":
        return extract_txt(path)